import os

def create_txt_files():
    # 1. Space Exploration History
    space_content = """# Key Milestones in Space Exploration: A Historical Overview

## Section 1: The Dawn of the Space Age
The exploration of space by humanity began in earnest in the mid-20th century. The space age was officially inaugurated on October 4, 1957, when the Soviet Union launched Sputnik 1, the first artificial satellite, into Earth orbit. Sputnik 1 was a small metal sphere with four external radio antennas that broadcast radio pulses back to Earth, proving that orbital insertion was achievable. This event triggered the Space Race between the United States and the Soviet Union. Shortly after, on April 12, 1961, Soviet cosmonaut Yuri Gagarin became the first human in space, orbiting the Earth in Vostok 1. His flight lasted 108 minutes and demonstrated that humans could survive in microgravity, launching a new epoch of human spaceflight.

## Section 2: The Apollo Moon Landings
In response to early Soviet successes, U.S. President John F. Kennedy declared in 1961 the ambitious goal of landing a man on the Moon and returning him safely to Earth before the decade was out. This goal was achieved by NASA's Apollo 11 mission on July 20, 1969. Commander Neil Armstrong and Lunar Module Pilot Buzz Aldrin landed the Apollo Lunar Module "Eagle" on the lunar surface in the Sea of Tranquility. Armstrong became the first person to walk on the Moon, famously declaring: "That's one small step for [a] man, one giant leap for mankind." The Apollo program ultimately landed twelve astronauts on the Moon across six successful missions, conducting experiments, taking photographs, and bringing back 382 kilograms (842 pounds) of lunar rocks and soil, which transformed our understanding of planetary science.

## Section 3: Space Shuttle and Space Stations
Following the Moon landings, space exploration shifted focus toward reusable transport and long-duration space habitation. In 1981, NASA launched the Space Shuttle Columbia, initiating the Space Shuttle program. The shuttle was the world's first reusable spacecraft, designed to carry heavy payloads into orbit and return like a glider. It was instrumental in launching the Hubble Space Telescope and building the International Space Station (ISS). The ISS, a collaborative project involving five space agencies (NASA, Roscosmos, ESA, JAXA, and CSA), has been continuously occupied by humans since November 2000. It serves as a microgravity and space environment research laboratory in low Earth orbit, where crew members conduct experiments in biology, physics, astronomy, and meteorology.

## Section 4: Robotic Mars Exploration
While humans stayed in low Earth orbit, robotic explorers ventured deeper into the solar system. Mars has been a primary target for robotic exploration. In 1997, NASA's Mars Pathfinder mission deployed Sojourner, the first successful rover on another planet. This was followed by the twin rovers Spirit and Opportunity in 2004, which found extensive evidence of past liquid water on Mars. In 2012, the car-sized Curiosity rover landed in Gale Crater to investigate Martian climate and geology, discovering organic molecules and evidence of ancient freshwater lakes. The Perseverance rover, landing in Jezero Crater in February 2021, represents the state of the art, tasked with searching for signs of ancient microbial life and collecting rock samples for future return to Earth.

## Section 5: The Artemis Program and Future Horizons
In the modern era, space exploration is entering a commercial and multilateral phase. NASA's Artemis program aims to land the first woman and the next man on the Moon, establishing a sustainable, long-term human presence on the lunar surface as a stepping stone to Mars. Artemis utilizes the Space Launch System (SLS), the most powerful rocket ever built, and the Orion spacecraft. In parallel, private commercial enterprises like SpaceX and Blue Origin are reducing launch costs through reusable rocket technology. The ultimate horizon remains human exploration of Mars, which will require advanced life support systems, radiation shielding, in-situ resource utilization, and nuclear thermal propulsion technologies.
"""
    os.makedirs("data", exist_ok=True)
    with open("data/space_exploration_history.txt", "w", encoding="utf-8") as f:
        f.write(space_content)
    print("Created data/space_exploration_history.txt")

    # 2. Mental Health and Sleep
    sleep_content = """# Sleep and Mental Health: A Neuroscientific Perspective

## Section 1: The Neuroscience of Sleep
Sleep is not a passive state of inactivity, but rather an active, highly regulated biological process essential for brain function and survival. Normal human sleep is organized into repeating 90-to-110-minute cycles, alternating between Non-Rapid Eye Movement (NREM) sleep and Rapid Eye Movement (REM) sleep. NREM sleep is divided into three stages: Stage 1 (light sleep), Stage 2 (deeper sleep with characteristic sleep spindles and K-complexes), and Stage 3 (slow-wave sleep or deep sleep). Slow-wave sleep is characterized by high-amplitude delta waves and is critical for physical recovery, tissue repair, and clearance of metabolic waste from the brain via the glymphatic system. REM sleep, characterized by rapid eye movements, temporary muscle paralysis, and vivid dreaming, plays a crucial role in memory consolidation, particularly emotional memory, and cognitive processing.

## Section 2: Circadian Rhythms and Biological Clocks
Every cell in the human body operates on a near-24-hour biological cycle known as a circadian rhythm. These rhythms are coordinated by a master biological clock located in the suprachiasmatic nucleus (SCN) of the hypothalamus in the brain. The SCN is highly sensitive to external light cues, which are detected by specialized light-sensitive ganglion cells in the retina and transmitted via the retinohypothalamic tract. In response to darkness, the SCN signals the pineal gland to synthesize and secrete the hormone melatonin, which promotes sleep. Exposure to artificial blue light (common in smartphones, tablets, and computer screens) in the evening suppresses melatonin secretion, shifting the circadian phase and making it difficult to fall asleep, eventually leading to chronic sleep deficits.

## Section 3: The Bidirectional Link
For decades, sleep disturbances were viewed merely as symptoms or side effects of psychiatric disorders like depression and anxiety. However, modern neuroscience has revealed a bidirectional, causal relationship. Sleep disruption can directly contribute to the onset and maintenance of mental health issues. During healthy REM sleep, the brain processes emotional experiences in a state free of the stress chemical noradrenaline, which acts as a form of nocturnal therapy. Deprivation of REM sleep prevents this emotional calibration, leading to hyper-reactivity in the amygdala—the brain's emotional threat center. Consequently, chronic sleep deprivation increases the risk of developing clinical depression, panic disorders, and generalized anxiety. Conversely, psychiatric conditions often disrupt neural circuits regulating sleep, creating a self-reinforcing pathological cycle.

## Section 4: Sleep Hygiene Best Practices
Improving sleep quality, known as sleep hygiene, is a powerful and accessible intervention for enhancing mental health and cognitive performance. Key practices include:
1. Maintaining a consistent sleep schedule: Going to bed and waking up at the same time every day, including weekends, stabilizes the circadian rhythm.
2. Optimizing the sleeping environment: The bedroom should be dark, quiet, and cool, ideally between 15 and 19 degrees Celsius (60 to 67 degrees Fahrenheit), to facilitate the natural drop in core body temperature required for sleep.
3. Avoiding blue-spectrum light: Turning off screens at least 60 minutes before bedtime reduces SCN stimulation.
4. Restricting stimulants and depressants: Caffeine blocks adenosine receptors (which signal sleep pressure) and should be avoided in the afternoon, while alcohol disrupts REM sleep architecture.
5. Engaging in regular exercise: Physical activity deepens slow-wave sleep, but intensive workouts should be completed several hours before bed.
"""
    with open("data/mental_health_and_sleep.txt", "w", encoding="utf-8") as f:
        f.write(sleep_content)
    print("Created data/mental_health_and_sleep.txt")

def create_docx_file():
    try:
        import docx
    except ImportError:
        print("python-docx is not installed yet. Skipping DOCX creation for a moment.")
        return

    doc = docx.Document()
    
    # Title
    doc.add_heading("Introduction to Quantum Computing", level=0)
    
    # Section 1
    doc.add_heading("Section 1: Core Principles of Quantum Mechanics", level=1)
    doc.add_paragraph(
        "Quantum computing is a revolutionary paradigm that leverages the laws of quantum mechanics to process "
        "information in ways that classical computers cannot. While classical computers represent information using "
        "bits, which can be either 0 or 1, quantum computers use quantum bits, or qubits. Qubits can exist in a state "
        "of superposition. Superposition allows a qubit to represent a 0, a 1, or any linear combination of both "
        "simultaneously. This means that a quantum system with N qubits can represent 2^N states at once, enabling "
        "massive parallel computation pathways."
    )
    doc.add_paragraph(
        "Another fundamental quantum principle is entanglement. When qubits become entangled, the state of one qubit "
        "instantaneously determines the state of another, no matter how far apart they are. Albert Einstein famously "
        "referred to this phenomenon as 'spooky action at a distance.' Entanglement allows quantum computers to "
        "coordinate calculations across qubits with incredible efficiency. Finally, quantum measurement collapse "
        "occurs when a qubit in superposition is observed; it instantly collapses into a definite classical state of "
        "either 0 or 1, dictated by the probability amplitudes of its superposition state."
    )

    # Section 2
    doc.add_heading("Section 2: Quantum Gates and Circuits", level=1)
    doc.add_paragraph(
        "To perform computations, quantum computers manipulate qubits using quantum gates. Unlike classical logic gates "
        "(such as AND, OR, NOT) which operate on deterministic bit inputs, quantum gates perform unitary mathematical "
        "rotations on the state vectors of qubits. For example, the Hadamard (H) gate is used to put a qubit into a "
        "state of equal superposition, converting a definite |0> state into a state where it has a 50% chance of being "
        "measured as 0 and a 50% chance of being measured as 1."
    )
    doc.add_paragraph(
        "Another critical gate is the Controlled-NOT (CNOT) gate, which operates on two qubits: a control qubit and a "
        "target qubit. The CNOT gate flips the state of the target qubit if and only if the control qubit is in the |1> "
        "state. The CNOT gate is essential for creating quantum entanglement between qubits. By chaining together "
        "single-qubit gates and multi-qubit gates in a sequence, designers construct quantum circuits that execute "
        "complex algorithms."
    )

    # Section 3
    doc.add_heading("Section 3: Quantum Algorithms", level=1)
    doc.add_paragraph(
        "The theoretical power of quantum computing is realized through specialized algorithms. The most famous of "
        "these is Shor's Algorithm, developed by Peter Shor in 1994. Shor's algorithm can factor large integers in "
        "polynomial time, which is exponentially faster than the best-known classical algorithms. Since modern encryption "
        "schemes like RSA rely on the extreme difficulty of integer factorization, Shor's algorithm represents a "
        "significant threat to global cybersecurity, prompting the development of post-quantum cryptography."
    )
    doc.add_paragraph(
        "Another milestone is Grover's Algorithm, formulated by Lov Grover in 1996. Grover's algorithm performs a "
        "quantum search on an unsorted database of N elements in O(sqrt(N)) time. This quadratic speedup is incredibly "
        "useful for optimization and search problems, reducing a search space of 1 million items to just 1,000 steps. "
        "Other algorithms target quantum chemistry and materials science simulations, allowing researchers to model "
        "molecular structures and chemical reactions at an atomic level with perfect accuracy."
    )

    # Section 4
    doc.add_heading("Section 4: Physical Implementations and Technical Challenges", level=1)
    doc.add_paragraph(
        "Building a physical quantum computer is one of the greatest engineering challenges of our time. Qubits are "
        "highly sensitive to their external environment. Any interaction with thermal, electromagnetic, or mechanical "
        "noise can cause quantum decoherence, where the qubits lose their quantum properties and collapse back into "
        "classical states. To prevent decoherence, many quantum architectures, such as superconducting qubits (used by "
        "IBM and Google), must be cooled in dilution refrigerators to temperatures near absolute zero, around 10 millikelvin, "
        "which is colder than deep space."
    )
    doc.add_paragraph(
        "Other physical approaches include trapped ion systems, which use electromagnetic fields to suspend individual "
        "ions in a vacuum and manipulate them with lasers, and topological qubits, which aim to store information "
        "braided in paths to inherently protect against noise. To handle unavoidable errors, researchers are developing "
        "Quantum Error Correction (QEC) schemes, which combine multiple physical qubits into a single logical qubit, "
        "though this requires a high overhead of physical qubits per logical qubit."
    )
    
    os.makedirs("data", exist_ok=True)
    doc.save("data/quantum_computing_intro.docx")
    print("Created data/quantum_computing_intro.docx")

def create_pdf_files():
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    except ImportError:
        print("reportlab is not installed yet. Skipping PDF creation for a moment.")
        return

    os.makedirs("data", exist_ok=True)

    # Helper function to generate PDF
    def build_pdf(filename, title, sections):
        doc = SimpleDocTemplate(f"data/{filename}", pagesize=letter,
                                rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        story = []
        styles = getSampleStyleSheet()

        # Custom paragraph style
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=10,
            leading=14,
            spaceAfter=10
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            leading=18,
            spaceBefore=15,
            spaceAfter=8,
            keepWithNext=True
        )

        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=22,
            leading=26,
            spaceAfter=20
        )

        # Title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 15))

        # Content
        for idx, (section_title, paragraphs) in enumerate(sections):
            if idx > 0 and idx % 2 == 0:
                # Add page break to ensure it spans at least 2 pages
                story.append(PageBreak())
            story.append(Paragraph(section_title, heading_style))
            for p_text in paragraphs:
                story.append(Paragraph(p_text, body_style))
            story.append(Spacer(1, 10))

        doc.build(story)
        print(f"Created data/{filename}")

    # 1. Artificial Intelligence Guide
    ai_sections = [
        ("Section 1: Foundations of Artificial Intelligence", [
            "Artificial Intelligence (AI) refers to the simulation of human intelligence in machines that are programmed to think like humans and mimic their actions. The term may also be applied to any machine that exhibits traits associated with a human mind such as learning and problem-solving. At its core, modern AI relies on Machine Learning (ML), a subset of AI that enables systems to learn from data, identify patterns, and make decisions with minimal human intervention.",
            "Within Machine Learning, Supervised Learning is the most common paradigm. In this approach, models are trained on labeled datasets, meaning each training example is paired with its correct output. The model learns a mapping function from input to output, allowing it to predict labels for new, unseen data. In contrast, Unsupervised Learning involves training models on unlabeled data, where the goal is to discover hidden structures, groupings, or distributions within the data itself (such as clustering customer segments)."
        ]),
        ("Section 2: Deep Learning and Neural Networks", [
            "Deep Learning is a specialized subfield of Machine Learning inspired by the structure and function of the human brain's neural networks. Artificial Neural Networks (ANNs) consist of interconnected nodes, or neurons, organized in layers: an input layer, one or more hidden layers, and an output layer. In a deep neural network, there are multiple hidden layers, which allow the model to learn increasingly abstract representations of the input data.",
            "During training, data flows forward through the network (forward propagation) to produce a prediction, which is compared to the actual label using a loss function. The error is then backpropagated through the network using gradient descent, adjusting the weights and biases of the connections to minimize the loss. Deep learning has driven breakthroughs in computer vision, speech recognition, and natural language processing, surpassing classical algorithms on high-dimensional data."
        ]),
        ("Section 3: The Transformer Architecture", [
            "Introduced in the seminal 2017 paper 'Attention Is All You Need' by Vaswani et al., the Transformer architecture revolutionized natural language processing (NLP). Prior to Transformers, recurrent neural networks (RNNs) and long short-term memory (LSTM) networks were the standard for sequential data. However, RNNs process text sequentially, token-by-token, which limits parallelization during training and makes it difficult to capture long-range dependencies.",
            "The Transformer replaces recurrence entirely with the Self-Attention mechanism. Self-attention allows the model to look at other tokens in a sentence to gain a better encoding for the current token, regardless of their distance. Multi-head attention projects the query, key, and value vectors into multiple subspaces, allowing the model to attend to information from different representation spaces simultaneously. This parallelized structure enables training on massive web-scale corpora, forming the foundation of modern Large Language Models (LLMs) like GPT and Gemini."
        ]),
        ("Section 4: Prompt Engineering Principles", [
            "Prompt Engineering is the practice of structuring text inputs (prompts) to guide Large Language Models to produce desired, high-quality outputs. Since LLMs generate text based on probabilistic patterns learned during pretraining, the phrasing and context of a prompt can dramatically influence the response. Effective prompt engineering utilizes specific techniques to improve model reasoning and factual accuracy.",
            "One common technique is Few-Shot Prompting, where the user provides a few examples of input-output pairs to demonstrate the desired format or task style before asking the actual question. Another powerful method is Chain-of-Thought (CoT) Prompting, which instructs the model to break down its reasoning step-by-step before arriving at a final answer. This is particularly effective for complex mathematical or logical reasoning tasks. System Prompts are also used to set the overall behavior, persona, constraints, and instructions for the model's entire session."
        ])
    ]
    build_pdf("artificial_intelligence_guide.pdf", "A Comprehensive Guide to Artificial Intelligence and Transformers", ai_sections)

    # 2. Renewable Energy Future
    energy_sections = [
        ("Section 1: Solar Energy Technologies", [
            "Solar energy is one of the most abundant and clean sources of renewable power available. The primary technology used to harvest solar energy is photovoltaics (PV), which directly converts sunlight into electricity using semiconductor materials. Traditional solar panels rely on silicon cells. However, recent scientific advances have introduced perovskite solar cells. Perovskite materials offer high absorption coefficients and can be processed at low temperatures, significantly reducing manufacturing costs while achieving high conversion efficiencies.",
            "Another approach is Concentrated Solar Power (CSP), which uses mirrors or lenses to focus a large area of sunlight onto a receiver. The intense heat generated is used to drive a steam turbine connected to an electrical generator. CSP systems often incorporate thermal energy storage, such as molten salt, allowing them to continue generating electricity even when the sun is not shining, providing valuable dispatchable power to the electrical grid."
        ]),
        ("Section 2: Wind Power and Turbine Design", [
            "Wind power utilizes the kinetic energy of flowing air to turn wind turbines and generate electricity. Wind farms can be located onshore (in open fields or plains) or offshore (in bodies of water). Offshore wind farms experience stronger and more consistent wind speeds, resulting in higher capacity factors. However, they also present greater engineering challenges, including corrosion from saltwater and the need for undersea transmission cables.",
            "Modern wind turbine designs have grown significantly in size and capacity. The power output of a wind turbine is proportional to the swept area of its blades, prompting manufacturers to build larger rotors and taller towers. Advanced materials, such as carbon-fiber reinforced polymers, are used to keep blades lightweight yet structurally sound. Gearless direct-drive generators are also becoming common, reducing mechanical wear and maintenance costs over the turbine's lifespan."
        ]),
        ("Section 3: Geothermal and Hydroelectric Systems", [
            "Geothermal energy taps into the heat stored beneath the Earth's crust. It relies on drilling deep wells to access hot water and steam from geothermal reservoirs, which is then piped to the surface to drive steam turbines. Geothermal power is highly reliable and provides continuous baseline electricity, unaffected by weather conditions. However, it is geographically limited to regions with active tectonic boundaries or high geothermal gradients.",
            "Hydroelectric power, which harnesses the energy of falling or flowing water, is the largest source of renewable electricity globally. Conventional hydroelectric dams impound water in a reservoir, releasing it through turbines to generate power. Run-of-river systems generate electricity without large reservoirs, minimizing environmental impacts but making output dependent on seasonal river flows. Pumped-storage hydropower acts as a giant water battery, pumping water to an upper reservoir during low-demand periods and releasing it through turbines during peak demand."
        ]),
        ("Section 4: Grid Integration and Energy Storage", [
            "Integrating variable renewable energy sources, like solar and wind, into the electrical grid presents significant challenges due to their intermittent nature. The grid must maintain a constant balance between electricity supply and demand to prevent frequency fluctuations and blackouts. When solar and wind generation exceeds demand, excess energy is lost (curtailed) unless it can be stored.",
            "Battery Energy Storage Systems (BESS), utilizing lithium-ion or flow batteries, are critical for buffering intermittency. They can store excess power and discharge it within milliseconds when supply drops. Additionally, Smart Grid technologies—incorporating advanced sensors, smart meters, and machine learning forecasting algorithms—are being deployed to dynamically balance loads, optimize power distribution, and coordinate distributed energy resources across the grid."
        ])
    ]
    build_pdf("renewable_energy_future.pdf", "The Future of Renewable Energy and Grid Integration", energy_sections)


if __name__ == "__main__":
    create_txt_files()
    
    # Try docx and pdf, they might fail if dependencies are not yet installed
    # but the user or system will run this script after pip install
    create_docx_file()
    create_pdf_files()
